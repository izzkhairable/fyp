from docx import Document
from docx.shared import Inches
import pyodbc
import json
import configparser
from docx.enum.text import WD_ALIGN_PARAGRAPH

config = configparser.ConfigParser()
config.read('sql_connect.cfg')

conn = pyodbc.connect('Driver=' + config['database']['driver'] + ';'
                      'Server=' + config['database']['server'] + ';'
                      'Database=' + config['database']['database'] + ';'
                      'Trusted_Connection=' + config['database']['trusted_connection'] + ';')
cursor = conn.cursor()


quotation = cursor.execute(
    "select labour_cost, labour_no_of_hours, markup_pct, testing_cost, remark from dbo.quotation where quotation_no=?", 'quotation_one')
res = quotation.fetchone()

markup_pct = None
if res.markup_pct != None:
    markup_pct = res.markup_pct/100


row_level = cursor.execute(
    "select id, row, component_no, description, quantity, uom, unit_price, remark, crawl_info, bom_id from dbo.quotation_component where quotation_no=? ORDER BY row ASC", 'quotation_one')
result = row_level.fetchall()

total_price = 0

for each_row in result:
    unit_price = 0
    qty = 0
    if each_row.unit_price != None:
        unit_price = each_row.unit_price
    if each_row.quantity != None:
        qty = each_row.quantity
    price = unit_price * qty
    total_price+=price
    

total_price = (total_price * markup_pct) + total_price + (res.labour_cost * res.labour_no_of_hours) + res.testing_cost

document = Document('aqs_bot/quotation_generation/Official Quotation Template.docx')

document.tables[3].cell(0,2).text = document.tables[3].cell(0,2).text.replace("{total_price}", str(total_price))
img = document.tables[3].cell(2,2).add_paragraph()
img.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = img.add_run()
r.add_picture('signature.jpg', width=Inches(0.8))


table = document.tables[2]
sn = 1
row = table.add_row().cells#https://github.com/python-openxml/python-docx/issues/205
row[0].text = str(sn)
row[1].text = "quotation_one"
row[3].text = "1"
row[4].text = "SET"
row[5].text = str(total_price)
row[6].text = str(total_price)


document.save('aqs_bot/quotation_generation/official_quote.docx')