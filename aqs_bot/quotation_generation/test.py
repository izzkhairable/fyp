from docx import Document
from docx.shared import Inches
import pyodbc
import json

conn = pyodbc.connect('Driver={SQL Server};'
                      'Server=DESKTOP-KMU57HS;'
                      'Database=myerp101;'
                      'Trusted_Connection=yes;')
cursor = conn.cursor()
# quotation = cursor.execute(
#     "select labour_cost, markup_pct, labour_cost_description from dbo.quotation where quotation_no=?", 'quotation_one')
# res = quotation.fetchone()

# markup_pct = None
# if res.markup_pct != None:
#     markup_pct = res.markup_pct/100

row_level = cursor.execute(
    "select row, component_no, description, quantity, uom, unit_price, remark, crawl_info from dbo.quotation_component where quotation_no=?", 'quotation_one')
result = row_level.fetchall()

for each_row in result:
    multiple_supplier = False

    if each_row.crawl_info != None:
        crawl_info = json.loads(each_row.crawl_info)
        if len(crawl_info) > 1:
            for ea_supplier in crawl_info:
                
            multiple_supplier = True

    if multiple_supplier == False:
        


 








# document = Document('aqs_bot/quotation_generation/Official Quotation Template.docx')

# print(document.tables[1].cell(0,1).text)

# table = document.tables[2]
# row = table.add_row().cells#https://github.com/python-openxml/python-docx/issues/205

# row[0].text = 'Hi'
# row[1].text = 'Bye'


# document.save('aqs_bot/quotation_generation/baby.docx')